"""
Script to generate Functional and Technical Documentation in .docx format
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

def setup_document_styles(doc):
    """Setup custom styles for the document"""
    styles = doc.styles
    
    # Heading 1 style
    h1_style = styles['Heading 1']
    h1_font = h1_style.font
    h1_font.name = 'Calibri'
    h1_font.size = Pt(18)
    h1_font.bold = True
    h1_font.color.rgb = RGBColor(26, 115, 232)
    
    # Heading 2 style
    h2_style = styles['Heading 2']
    h2_font = h2_style.font
    h2_font.name = 'Calibri'
    h2_font.size = Pt(14)
    h2_font.bold = True
    h2_font.color.rgb = RGBColor(21, 87, 176)
    
    # Heading 3 style
    h3_style = styles['Heading 3']
    h3_font = h3_style.font
    h3_font.name = 'Calibri'
    h3_font.size = Pt(12)
    h3_font.bold = True
    
    # Normal style
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)

    # Other styles to ensure consistency
    for style_name in ['No Spacing', 'List Paragraph', 'List Bullet', 'List Number', 'List Bullet 2']:
        try:
            if style_name in styles:
                style = styles[style_name]
                if hasattr(style, 'font'):
                    style.font.name = 'Calibri'
                    style.font.size = Pt(11)
        except Exception:
            pass  # Ignore if style modifications fail

def add_header(doc, title, subtitle):
    """Add document header"""
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    title_run = header_para.add_run(title)
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(26, 115, 232)
    
    doc.add_paragraph()
    
    meta_para = doc.add_paragraph()
    meta_run = meta_para.add_run(subtitle)
    meta_run.font.name = 'Calibri'
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(102, 102, 102)
    
    doc.add_paragraph()

def add_table_of_contents(doc, items):
    """Add table of contents"""
    doc.add_heading('Table of Contents', level=1)
    doc.add_paragraph()
    
    for i, item in enumerate(items, 1):
        para = doc.add_paragraph(f"{i}. {item}", style='List Number')
        para.style.font.name = 'Calibri'
        para.style.font.size = Pt(11)

def create_functional_documentation():
    """Create Functional Documentation in .docx format"""
    doc = Document()
    setup_document_styles(doc)
    
    # Header
    add_header(doc, 'Functional Documentation', 
               f'Application: Google Patent PDF Downloader\nVersion: 1.0\nDocument Version: 1.0\nDate: December 2025')
    
    # Table of Contents
    toc_items = [
        'Overview',
        'Purpose and Scope',
        'User Requirements',
        'Functional Features',
        'User Interface',
        'User Workflows',
        'Input Requirements',
        'Output Specifications',
        'Error Handling',
        'Use Cases',
        'Business Rules',
        'Non-Functional Requirements'
    ]
    add_table_of_contents(doc, toc_items)
    
    doc.add_page_break()
    
    # Section 1: Overview
    doc.add_heading('1. Overview', level=1)
    doc.add_paragraph(
        'The Patent Downloader is a desktop application designed to automate the process of downloading patent documents from Google Patents. '
        'It provides a user-friendly graphical interface that allows users to batch download multiple patent PDFs by reading patent numbers from an Excel file.'
    )
    
    doc.add_heading('Key Benefits', level=2)
    benefits = [
        ('Time Savings', 'Automates manual patent downloading process'),
        ('Batch Processing', 'Download hundreds of patents in a single operation'),
        ('User-Friendly', 'Intuitive GUI requiring minimal technical knowledge'),
        ('Reliable', 'Multiple fallback methods ensure high success rate'),
        ('Organized', 'Automatic file naming and organization'),
        ('Transparent', 'Detailed logging and progress tracking')
    ]
    for title, desc in benefits:
        para = doc.add_paragraph()
        run1 = para.add_run(f'{title}: ')
        run1.bold = True
        para.add_run(desc)
    
    # Section 2: Purpose and Scope
    doc.add_page_break()
    doc.add_heading('2. Purpose and Scope', level=1)
    
    doc.add_heading('2.1 Purpose', level=2)
    doc.add_paragraph(
        'The application enables patent researchers, legal professionals, and analysts to efficiently download patent documents in PDF format '
        'from Google Patents without manual intervention.'
    )
    
    doc.add_heading('2.2 Scope', level=2)
    
    doc.add_heading('In Scope:', level=3)
    in_scope = [
        'Reading patent numbers from Excel files',
        'Downloading patent PDFs from Google Patents',
        'Fallback to FreePatentsOnline when Google Patents fails',
        'Real-time progress tracking',
        'Logging of successful and failed downloads',
        'Generating Excel reports with patent metadata'
    ]
    for item in in_scope:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Out of Scope:', level=3)
    out_scope = [
        'Patent search functionality',
        'Patent analysis features',
        'Integration with other patent databases',
        'Patent document editing or annotation',
        'Cloud storage integration'
    ]
    for item in out_scope:
        doc.add_paragraph(item, style='List Bullet')
    
    # Section 3: User Requirements
    doc.add_page_break()
    doc.add_heading('3. User Requirements', level=1)
    
    doc.add_heading('3.1 Target Users', level=2)
    users = [
        ('Patent Researchers', 'Need to download multiple patents for analysis'),
        ('Legal Professionals', 'Require patent documents for legal proceedings'),
        ('R&D Teams', 'Collect patents for competitive intelligence'),
        ('Academic Researchers', 'Download patents for research purposes')
    ]
    for i, (user, desc) in enumerate(users, 1):
        para = doc.add_paragraph()
        run1 = para.add_run(f'{i}. {user}: ')
        run1.bold = True
        para.add_run(desc)
    
    doc.add_heading('3.2 User Skills Required', level=2)
    skills = [
        'Basic computer literacy',
        'Ability to use Excel files',
        'No programming knowledge required'
    ]
    for skill in skills:
        doc.add_paragraph(skill, style='List Bullet')
    
    doc.add_heading('3.3 System Requirements', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Requirement'
    hdr_cells[1].text = 'Specification'
    
    requirements = [
        ('Operating System', 'Windows 10 or higher'),
        ('Python', 'Version 3.8 or higher'),
        ('Browser', 'Google Chrome installed'),
        ('Internet Connection', 'Required for downloading patents'),
        ('Excel File', 'Must contain patent numbers in "Display Key" column')
    ]
    for req, spec in requirements:
        row_cells = table.add_row().cells
        row_cells[0].text = req
        row_cells[1].text = spec
    
    # Section 4: Functional Features
    doc.add_page_break()
    doc.add_heading('4. Functional Features', level=1)
    
    doc.add_heading('4.1 Core Features', level=2)
    
    doc.add_heading('4.1.1 Excel File Reading', level=3)
    doc.add_paragraph('Description: Reads patent numbers from Excel files (.xlsx, .xls)')
    doc.add_paragraph('Input: Excel file with "Display Key" column')
    doc.add_paragraph('Output: List of patent numbers extracted from the file')
    doc.add_paragraph('Validation: Checks for column existence and validates file format')
    
    doc.add_heading('4.1.2 Patent Download', level=3)
    doc.add_paragraph('Description: Downloads patent PDFs from Google Patents')
    doc.add_paragraph('Primary Method: Direct HTTP download (faster, no browser)')
    doc.add_paragraph('Fallback Method: Browser automation when direct download fails')
    doc.add_paragraph('Secondary Fallback: FreePatentsOnline as alternative source')
    doc.add_paragraph('Output: PDF files saved with patent numbers as filenames')
    
    doc.add_heading('4.1.3 Progress Tracking', level=3)
    doc.add_paragraph('Description: Real-time display of download progress')
    doc.add_paragraph('Components:')
    components = [
        'Progress bar showing percentage completion',
        'Status messages indicating current operation',
        'Activity log showing detailed operations',
        'Success/failure counts'
    ]
    for comp in components:
        doc.add_paragraph(comp, style='List Bullet 2')
    
    doc.add_heading('4.1.4 File Management', level=3)
    doc.add_paragraph('Description: Organizes downloaded PDFs')
    doc.add_paragraph('Features:')
    features = [
        'Automatic folder creation (downloaded_patents/)',
        'Consistent file naming (PatentNumber.pdf)',
        'One-click access to downloads folder'
    ]
    for feat in features:
        doc.add_paragraph(feat, style='List Bullet 2')
    
    doc.add_heading('4.2 Advanced Features', level=2)
    
    doc.add_heading('4.2.1 Logging System', level=3)
    doc.add_paragraph('Main Log: Records all operations, successes, and errors')
    doc.add_paragraph('Failed Patents Log: Dedicated log for failed downloads with:')
    log_items = [
        'Original patent number',
        'Cleaned patent number',
        'Failure reason',
        'Patent URL',
        'Timestamp'
    ]
    for item in log_items:
        doc.add_paragraph(item, style='List Bullet 2')
    
    doc.add_heading('4.2.2 Excel Report Generation', level=3)
    doc.add_paragraph('Description: Creates Excel report with patent metadata')
    doc.add_paragraph('Contents:')
    contents = [
        'Patent Number',
        'Title',
        'Publication Date',
        'Applicant/Assignee',
        'Download Status',
        'Download Date'
    ]
    for content in contents:
        doc.add_paragraph(content, style='List Bullet 2')
    doc.add_paragraph('Update Frequency: Real-time updates after each successful download')
    
    # Section 5: User Interface
    doc.add_page_break()
    doc.add_heading('5. User Interface', level=1)
    
    doc.add_heading('5.1 Main Window Components', level=2)
    
    doc.add_heading('5.1.1 Header Section', level=3)
    doc.add_paragraph('Title: "📄 Patent Downloader"')
    doc.add_paragraph('Subtitle: "Download patents automatically from Google Patents"')
    doc.add_paragraph('Style: Dark header with white text')
    
    doc.add_heading('5.1.2 File Selection Section', level=3)
    doc.add_paragraph('File Entry Field: Displays selected Excel file path (read-only)')
    doc.add_paragraph('Browse Button: Opens file dialog to select Excel file')
    doc.add_paragraph('Info Label: Reminds user about "Display Key" column requirement')
    
    doc.add_heading('5.1.3 Download Control Section', level=3)
    doc.add_paragraph('Start Download Button: Initiates the download process')
    doc.add_paragraph('Stop Download Button: Stops ongoing download (disabled when not downloading)')
    
    doc.add_heading('5.1.4 Progress Section', level=3)
    doc.add_paragraph('Progress Bar: Visual indicator of download completion percentage')
    doc.add_paragraph('Status Label: Current operation status with icons')
    doc.add_paragraph('Activity Log: Scrollable text area showing detailed operations')
    
    doc.add_heading('5.1.5 Action Buttons (Bottom)', level=3)
    doc.add_paragraph('Open Downloads Folder: Opens file explorer to downloads directory')
    doc.add_paragraph('Main Log: Opens main activity log file')
    doc.add_paragraph('Failed Patents Log: Opens failed patents log file')
    
    doc.add_heading('5.2 Visual Design', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Element'
    hdr_cells[1].text = 'Color'
    hdr_cells[2].text = 'Hex Code'
    
    colors = [
        ('Primary Color', 'Blue', '#1a73e8'),
        ('Success Color', 'Green', '#34a853'),
        ('Error Color', 'Red', '#ea4335'),
        ('Accent Color', 'Yellow/Orange', '#fbbc04'),
        ('Background', 'Light Gray', '#f8f9fa')
    ]
    for elem, color, hex_code in colors:
        row_cells = table.add_row().cells
        row_cells[0].text = elem
        row_cells[1].text = color
        row_cells[2].text = hex_code
    
    # Section 6: User Workflows
    doc.add_page_break()
    doc.add_heading('6. User Workflows', level=1)
    
    doc.add_heading('6.1 Primary Workflow: Download Patents', level=2)
    
    steps = [
        ('Step 1: Launch Application', [
            'User double-clicks 🚀 START HERE - GUI.bat or runs python patent_downloader_gui.py',
            'Application window opens'
        ]),
        ('Step 2: Select Excel File', [
            'User clicks "📂 Browse Files" button',
            'File dialog opens',
            'User selects Excel file containing patent numbers',
            'File path appears in entry field'
        ]),
        ('Step 3: Start Download', [
            'User clicks "🚀 Start Download" button',
            'Application validates file and reads patent numbers',
            'Download process begins'
        ]),
        ('Step 4: Monitor Progress', [
            'Progress bar updates in real-time',
            'Activity log shows each patent being processed',
            'Status label shows current operation'
        ]),
        ('Step 5: Completion', [
            'Download completes automatically',
            'Summary dialog shows total patents processed, success count, failure count, location of downloaded files, Excel report location, and failed patents log location (if applicable)'
        ]),
        ('Step 6: Access Results', [
            'User can click "📂 Open Downloads Folder" to view PDFs',
            'User can click "📄 Main Log" to review operations',
            'User can click "❌ Failed Patents Log" to see failures'
        ])
    ]
    
    for step_title, step_items in steps:
        doc.add_heading(step_title, level=3)
        for item in step_items:
            doc.add_paragraph(item, style='List Bullet')
    
    # Section 7: Input Requirements
    doc.add_page_break()
    doc.add_heading('7. Input Requirements', level=1)
    
    doc.add_heading('7.1 Excel File Format', level=2)
    
    doc.add_heading('Required Structure:', level=3)
    doc.add_paragraph('File Format: .xlsx or .xls')
    doc.add_paragraph('Required Column: "Display Key" (case-sensitive)')
    doc.add_paragraph('Column Content: Patent numbers (one per row)')
    
    doc.add_heading('Example:', level=3)
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Display Key'
    
    examples = ['US1234567A', 'EP9876543B1', 'WO2020123456A1']
    for ex in examples:
        row_cells = table.add_row().cells
        row_cells[0].text = ex
    
    doc.add_heading('Patent Number Formats Supported:', level=3)
    formats = [
        'US patents: US1234567A, US12345678B2',
        'European patents: EP1234567A1, EP1234567B1',
        'PCT patents: WO2020123456A1',
        'Other formats: Any format recognized by Google Patents'
    ]
    for fmt in formats:
        doc.add_paragraph(fmt, style='List Bullet')
    
    # Section 8: Output Specifications
    doc.add_page_break()
    doc.add_heading('8. Output Specifications', level=1)
    
    doc.add_heading('8.1 Downloaded PDFs', level=2)
    doc.add_paragraph('Location: downloaded_patents/ folder (created automatically)')
    doc.add_paragraph('Naming Convention: Format: {PatentNumber}.pdf, Example: US1234567A.pdf')
    doc.add_paragraph('File Properties: Format: PDF (Portable Document Format), Source: Google Patents or FreePatentsOnline, Quality: Original patent document quality')
    
    doc.add_heading('8.2 Log Files', level=2)
    
    doc.add_heading('8.2.1 Main Log (patent_download_gui.log)', level=3)
    doc.add_paragraph('Format: Text file with timestamped entries')
    doc.add_paragraph('Content: Session start/end times, Excel file loading information, each patent download attempt, success/failure status, error messages, summary statistics')
    
    code_para = doc.add_paragraph('Example Entry:')
    code_para.style.font.name = 'Calibri'
    code_para.style.font.size = Pt(9)
    doc.add_paragraph('2025-12-15 10:30:45 - INFO - Selected file: C:\\patents.xlsx', style='No Spacing')
    doc.add_paragraph('2025-12-15 10:30:46 - INFO - Found 50 patent numbers', style='No Spacing')
    doc.add_paragraph('2025-12-15 10:30:47 - INFO - Downloading: US1234567A', style='No Spacing')
    doc.add_paragraph('2025-12-15 10:30:50 - INFO - SUCCESS', style='No Spacing')
    
    doc.add_heading('8.2.2 Failed Patents Log (failed_patents.log)', level=3)
    doc.add_paragraph('Format: Text file with structured entries')
    doc.add_paragraph('Content: Session separator, for each failed patent: Original patent number, Cleaned patent number, Failure reason, Patent URL, Timestamp')
    
    # Section 9: Error Handling
    doc.add_page_break()
    doc.add_heading('9. Error Handling', level=1)
    
    doc.add_heading('9.1 Error Categories', level=2)
    
    doc.add_heading('9.1.1 File Errors', level=3)
    doc.add_paragraph('Excel file not found: Error dialog displayed')
    doc.add_paragraph('Invalid Excel format: Error message in log')
    doc.add_paragraph('Missing "Display Key" column: Error message with available columns listed')
    
    doc.add_heading('9.1.2 Network Errors', level=3)
    doc.add_paragraph('Connection timeout: Automatic retry with fallback methods')
    doc.add_paragraph('Invalid patent number: Logged to failed_patents.log')
    doc.add_paragraph('PDF not found: Tries alternative sources, logs if all fail')
    
    doc.add_heading('9.1.3 System Errors', level=3)
    doc.add_paragraph('Chrome browser not available: Error dialog with installation instructions')
    doc.add_paragraph('Missing Python packages: Error message with installation command')
    doc.add_paragraph('Permission errors: Error dialog explaining folder access issues')
    
    # Section 10: Use Cases
    doc.add_page_break()
    doc.add_heading('10. Use Cases', level=1)
    
    use_cases = [
        {
            'title': 'Use Case 1: Batch Download for Research Project',
            'actor': 'Patent Researcher',
            'goal': 'Download 200 patents for competitive analysis',
            'steps': [
                'Researcher launches application',
                'Selects Excel file with 200 patent numbers',
                'Starts download',
                'Monitors progress (takes ~10-15 minutes)',
                'Reviews downloaded PDFs',
                'Checks failed_patents.log for any failures',
                'Reviews Excel report for metadata'
            ],
            'postcondition': '200 PDF files downloaded, Excel report generated'
        },
        {
            'title': 'Use Case 2: Quick Patent Download',
            'actor': 'Legal Professional',
            'goal': 'Download 5 specific patents for a case',
            'steps': [
                'Creates Excel file with 5 patent numbers',
                'Launches application',
                'Selects Excel file',
                'Starts download',
                'Downloads complete in ~1 minute',
                'Accesses PDFs from downloads folder'
            ],
            'postcondition': '5 PDF files ready for review'
        }
    ]
    
    for uc in use_cases:
        doc.add_heading(uc['title'], level=2)
        doc.add_paragraph(f"Actor: {uc['actor']}")
        doc.add_paragraph(f"Goal: {uc['goal']}")
        doc.add_paragraph('Steps:')
        for step in uc['steps']:
            doc.add_paragraph(step, style='List Number')
        doc.add_paragraph(f"Postcondition: {uc['postcondition']}")
        doc.add_paragraph()
    
    # Section 11: Business Rules
    doc.add_page_break()
    doc.add_heading('11. Business Rules', level=1)
    
    doc.add_heading('11.1 Download Rules', level=2)
    rules = [
        'One PDF per Patent: Each patent number results in one PDF file',
        'File Naming: PDFs named using cleaned patent number',
        'No Duplicates: Same patent number overwrites existing file',
        'Sequential Processing: Patents processed one at a time',
        'Delay Between Downloads: 2-second delay to avoid rate limiting'
    ]
    for rule in rules:
        para = doc.add_paragraph()
        parts = rule.split(':')
        run1 = para.add_run(parts[0] + ': ')
        run1.bold = True
        if len(parts) > 1:
            para.add_run(parts[1])
    
    doc.add_heading('11.2 Source Priority', level=2)
    priorities = [
        'Primary: Google Patents (direct download)',
        'Secondary: Google Patents (browser method)',
        'Tertiary: FreePatentsOnline'
    ]
    for priority in priorities:
        doc.add_paragraph(priority, style='List Number')
    
    # Section 12: Non-Functional Requirements
    doc.add_page_break()
    doc.add_heading('12. Non-Functional Requirements', level=1)
    
    nfr_sections = [
        ('12.1 Performance', [
            'Download Speed: 3-5x faster than manual download',
            'Direct Download Mode: No browser overhead for faster processing',
            'Memory Usage: Efficient memory management for large batches',
            'Response Time: GUI remains responsive during downloads'
        ]),
        ('12.2 Reliability', [
            'Success Rate: High success rate with multiple fallback methods',
            'Error Recovery: Graceful handling of network and system errors',
            'Data Integrity: Downloaded PDFs verified for completeness',
            'Session Recovery: Ability to resume interrupted downloads'
        ]),
        ('12.3 Usability', [
            'Ease of Use: No technical knowledge required',
            'Clear Interface: Intuitive GUI with clear labels',
            'Feedback: Real-time progress and status updates',
            'Documentation: Clear instructions and error messages'
        ]),
        ('12.4 Compatibility', [
            'Python Versions: Compatible with Python 3.8 through 3.14+',
            'Windows Versions: Windows 10 and higher',
            'Excel Formats: Supports .xlsx and .xls files',
            'Browser: Requires Google Chrome'
        ]),
        ('12.5 Security', [
            'No Data Collection: Application doesn\'t collect user data',
            'Local Processing: All operations performed locally',
            'Safe Downloads: Downloads from trusted sources only',
            'File Permissions: Respects system file permissions'
        ])
    ]
    
    for section_title, items in nfr_sections:
        doc.add_heading(section_title, level=2)
        for item in items:
            para = doc.add_paragraph()
            parts = item.split(':')
            run1 = para.add_run(parts[0] + ': ')
            run1.bold = True
            if len(parts) > 1:
                para.add_run(parts[1])

    # Section 13: Future Recommendations
    doc.add_page_break()
    doc.add_heading('13. Future Recommendations', level=1)
    
    doc.add_heading('13.1 Functional Enhancements', level=2)
    recs = [
        ('Cloud Integration', 'Support for direct upload to Google Drive, Dropbox, or OneDrive'),
        ('Email Notifications', 'Alert system to notify users via email when large batch processes complete'),
        ('Smart Resume', 'Ability to resume interrupted downloads exactly where they left off without rescanning'),
        ('Multi-Format Input', 'Support for CSV, TXT, and direct text entry of patent numbers in addition to Excel')
    ]
    for title, desc in recs:
        para = doc.add_paragraph()
        run1 = para.add_run(f'{title}: ')
        run1.bold = True
        para.add_run(desc)

    doc.add_heading('13.2 UI/UX Improvements', level=2)
    ui_recs = [
        'Dark Mode toggle for better accessibility',
        'Estimated time remaining calculation based on download speed',
        'Drag-and-drop support for file selection',
        'Preview pane for downloaded PDFs within the application'
    ]
    for item in ui_recs:
        doc.add_paragraph(item, style='List Bullet')

    # Appendices
    doc.add_page_break()
    doc.add_heading('Appendix A: Glossary', level=1)
    
    glossary_table = doc.add_table(rows=1, cols=2)
    glossary_table.style = 'Light Grid Accent 1'
    hdr_cells = glossary_table.rows[0].cells
    hdr_cells[0].text = 'Term'
    hdr_cells[1].text = 'Definition'
    
    glossary_terms = [
        ('Display Key', 'Column name in Excel file containing patent numbers'),
        ('Direct Download', 'Downloading PDF without opening browser'),
        ('Fallback Method', 'Alternative approach when primary method fails'),
        ('Patent Number', 'Unique identifier for a patent document'),
        ('PDF', 'Portable Document Format file'),
        ('ChromeDriver', 'Tool for automating Chrome browser')
    ]
    
    for term, definition in glossary_terms:
        row_cells = glossary_table.add_row().cells
        row_cells[0].text = term
        row_cells[1].text = definition
    
    doc.add_page_break()
    doc.add_heading('Appendix B: Version History', level=1)
    doc.add_paragraph('Version 1.0: Initial release including direct download mode, Excel reports, and logging system.')
    
    # Footer
    doc.add_paragraph()
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run('End of Functional Documentation')
    run.bold = True
    run.font.size = Pt(12)
    
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run('Google Patent PDF Downloader - Version 1.0')
    
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run('Document Version 1.0 - December 2025')
    
    # Save document
    doc.save('FUNCTIONAL_DOCUMENTATION.docx')
    print("Functional Documentation created: FUNCTIONAL_DOCUMENTATION.docx")

def create_technical_documentation():
    """Create Technical Documentation in .docx format"""
    doc = Document()
    setup_document_styles(doc)
    
    # Header
    add_header(doc, 'Technical Documentation',
               f'Application: Google Patent PDF Downloader\nVersion: 1.0\nDocument Version: 1.0\nDate: December 2025\nTechnology Stack: Python 3.8+, Tkinter, Selenium, Requests, BeautifulSoup4')
    
    # Table of Contents
    toc_items = [
        'System Architecture',
        'Technology Stack',
        'Code Structure',
        'Core Components',
        'Data Flow',
        'API and Methods',
        'Configuration',
        'Dependencies',
        'File Structure',
        'Implementation Details',
        'Error Handling Architecture',
        'Logging System',
        'Testing Considerations',
        'Deployment',
        'Performance Optimization'
    ]
    add_table_of_contents(doc, toc_items)
    
    doc.add_page_break()
    
    # Section 1: System Architecture
    doc.add_heading('1. System Architecture', level=1)
    
    doc.add_heading('1.1 High-Level Architecture', level=2)
    doc.add_paragraph(
        'The system follows a layered architecture with four main layers:'
    )
    layers = [
        ('Presentation Layer', 'Tkinter GUI handles user interaction'),
        ('Application Layer', 'Business logic orchestrates downloads'),
        ('Service Layer', 'Download methods interact with external APIs'),
        ('Data Layer', 'File system operations for PDFs and logs')
    ]
    for layer, desc in layers:
        para = doc.add_paragraph()
        run1 = para.add_run(f'{layer}: ')
        run1.bold = True
        para.add_run(desc)
    
    doc.add_heading('1.2 Component Interaction', level=2)
    doc.add_paragraph(
        'The system follows a layered architecture where each layer communicates with adjacent layers through well-defined interfaces.'
    )
    
    doc.add_heading('1.3 Threading Model', level=2)
    doc.add_paragraph('Main Thread: GUI event loop (Tkinter mainloop)')
    doc.add_paragraph('Download Thread: Separate daemon thread for download operations')
    doc.add_paragraph('Thread Safety: GUI updates via update_idletasks() from download thread')
    
    # Section 2: Technology Stack
    doc.add_page_break()
    doc.add_heading('2. Technology Stack', level=1)
    
    doc.add_heading('2.1 Core Technologies', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Technology'
    hdr_cells[1].text = 'Version'
    hdr_cells[2].text = 'Purpose'
    
    technologies = [
        ('Python', '3.8+', 'Core programming language'),
        ('Tkinter', 'Built-in', 'GUI framework'),
        ('Selenium', '4.15.0+', 'Browser automation (fallback)'),
        ('Requests', '2.31.0+', 'HTTP client for direct downloads'),
        ('BeautifulSoup4', '4.12.0+', 'HTML parsing for PDF URL extraction'),
        ('Pandas', '2.0.0+', 'Excel file reading'),
        ('OpenPyXL', '3.1.0+', 'Excel file writing')
    ]
    
    for tech, version, purpose in technologies:
        row_cells = table.add_row().cells
        row_cells[0].text = tech
        row_cells[1].text = version
        row_cells[2].text = purpose
    
    doc.add_heading('2.2 System Requirements', level=2)
    doc.add_heading('Runtime:', level=3)
    runtime_reqs = [
        'Python 3.8 or higher',
        'Google Chrome browser',
        'Windows 10 or higher',
        'Internet connection'
    ]
    for req in runtime_reqs:
        doc.add_paragraph(req, style='List Bullet')
    
    # Section 3: Code Structure
    doc.add_page_break()
    doc.add_heading('3. Code Structure', level=1)
    
    doc.add_heading('3.1 Main Module: patent_downloader_gui.py', level=2)
    doc.add_paragraph('Class Structure:')
    
    class_structure = [
        '__init__() - Initialization',
        'create_widgets() - GUI creation',
        'browse_file() - File selection',
        'start_download() - Download initiation',
        'stop_download() - Download termination',
        'download_patents() - Main download loop',
        'read_patent_numbers() - Excel reading',
        'download_patent() - Single patent download',
        'try_direct_download() - Direct HTTP download',
        'try_freepatentsonline() - FPO fallback',
        'download_pdf_direct() - PDF file download',
        'extract_patent_info() - Metadata extraction',
        'clean_patent_number() - Number normalization',
        'log_failed_patent() - Failure logging',
        'create_excel_report() - Excel report creation',
        'update_excel_report() - Excel report update'
    ]
    
    for method in class_structure:
        doc.add_paragraph(method, style='List Bullet')
    
    doc.add_heading('3.2 Utility Module: check_setup.py', level=2)
    doc.add_paragraph('Functions:')
    functions = [
        'check_python_version() - Validates Python version',
        'check_packages() - Verifies required packages',
        'check_chrome() - Tests Chrome/ChromeDriver availability'
    ]
    for func in functions:
        doc.add_paragraph(func, style='List Bullet')
    
    # Section 4: Core Components
    doc.add_page_break()
    doc.add_heading('4. Core Components', level=1)
    
    doc.add_heading('4.1 PatentDownloaderGUI Class', level=2)
    doc.add_paragraph('Purpose: Main application class managing GUI and download operations')
    
    doc.add_heading('Key Attributes:', level=3)
    attributes = [
        'self.root - Tkinter root window',
        'self.excel_file - Selected Excel file path',
        'self.output_dir - Downloads directory',
        'self.is_downloading - Download state flag',
        'self.driver - Selenium WebDriver instance',
        'self.failed_patents - List of failed patent info',
        'self.patent_info_list - Patent metadata for Excel',
        'self.colors - UI color scheme dictionary'
    ]
    for attr in attributes:
        doc.add_paragraph(attr, style='List Bullet')
    
    doc.add_heading('Key Methods:', level=3)
    
    methods = [
        ('__init__(self, root)', 'Initializes GUI window, sets up color scheme, creates output directory, initializes variables, calls create_widgets()'),
        ('create_widgets(self)', 'Builds complete GUI layout, creates all UI components, sets up event handlers, configures styling'),
        ('download_patents(self)', 'Main download orchestration method, reads patent numbers from Excel, iterates through patents, calls download methods, updates progress and logs, generates summary'),
        ('download_patent(self, patent_number)', 'Coordinates download attempts, tries direct download first, falls back to FreePatentsOnline, logs failures, returns success status')
    ]
    
    for method_name, method_desc in methods:
        para = doc.add_paragraph()
        run1 = para.add_run(f'{method_name}: ')
        run1.bold = True
        run1.font.name = 'Calibri'
        para.add_run(method_desc)
    
    doc.add_heading('4.2 Download Methods', level=2)
    
    doc.add_heading('try_direct_download(self, patent_number)', level=3)
    doc.add_paragraph('Purpose: Attempt direct PDF download without browser')
    doc.add_paragraph('Process:')
    process_steps = [
        'Construct Google Patents URL',
        'Fetch HTML using Requests',
        'Extract PDF URL using regex',
        'Download PDF directly',
        'Extract patent metadata',
        'Return success status'
    ]
    for step in process_steps:
        doc.add_paragraph(step, style='List Number')
    
    # Section 5: Data Flow
    doc.add_page_break()
    doc.add_heading('5. Data Flow', level=1)
    
    doc.add_heading('5.1 Download Flow', level=2)
    flow_steps = [
        'User Action: Start Download',
        'Read Excel File (pandas)',
        'Extract Patent Numbers',
        'For Each Patent:',
        '  - Try Direct Download',
        '    - Fetch HTML (requests)',
        '    - Extract PDF URL (regex)',
        '    - Download PDF (requests)',
        '    - Extract Metadata (BeautifulSoup)',
        '  - If Failed: Try FreePatentsOnline',
        '    - Construct FPO URL',
        '    - Download PDF (requests)',
        '  - If Failed: Log to failed_patents.log',
        'Update Progress Bar',
        'Update Excel Report',
        'Next Patent'
    ]
    
    for step in flow_steps:
        doc.add_paragraph(step, style='List Bullet')
    
    # Section 6: API and Methods
    doc.add_page_break()
    doc.add_heading('6. API and Methods', level=1)
    
    doc.add_heading('6.1 Public Methods', level=2)
    
    doc.add_heading('GUI Control Methods', level=3)
    gui_methods = [
        ('browse_file()', 'Opens file dialog, sets excel_file variable, updates log display'),
        ('start_download()', 'Validates file selection, starts download thread, updates button states'),
        ('stop_download()', 'Sets is_downloading flag to False, allows graceful termination'),
        ('open_output_folder()', 'Opens downloads folder in file explorer, uses os.startfile() on Windows')
    ]
    
    for method_name, method_desc in gui_methods:
        para = doc.add_paragraph()
        run1 = para.add_run(f'{method_name}: ')
        run1.bold = True
        run1.font.name = 'Calibri'
        para.add_run(method_desc)
    
    # Section 7: Configuration
    doc.add_page_break()
    doc.add_heading('7. Configuration', level=1)
    
    doc.add_heading('7.1 Application Configuration', level=2)
    
    doc.add_heading('Color Scheme:', level=3)
    color_code = doc.add_paragraph()
    color_code.style.font.name = 'Calibri'
    color_code.style.font.size = Pt(9)
    color_code.add_run('colors = {\n')
    color_code.add_run("    'primary': '#1a73e8',\n")
    color_code.add_run("    'primary_dark': '#1557b0',\n")
    color_code.add_run("    'success': '#34a853',\n")
    color_code.add_run("    'danger': '#ea4335',\n")
    color_code.add_run("    'background': '#f8f9fa',\n")
    color_code.add_run("    'surface': '#ffffff'\n")
    color_code.add_run('}')
    
    doc.add_heading('Window Configuration:', level=3)
    doc.add_paragraph('Default size: 1000x850')
    doc.add_paragraph('Minimum size: 900x750')
    doc.add_paragraph('Resizable: True')
    doc.add_paragraph('Title: "Google Patent PDF Downloader"')
    
    doc.add_heading('Download Configuration:', level=3)
    doc.add_paragraph('Output directory: downloaded_patents')
    doc.add_paragraph('Delay between downloads: 2 seconds')
    doc.add_paragraph('Timeout: 10 seconds (direct), 15 seconds (FPO)')
    doc.add_paragraph('Chunk size: 8192 bytes')
    
    # Section 8: Dependencies
    doc.add_page_break()
    doc.add_heading('8. Dependencies', level=1)
    
    doc.add_heading('8.1 Required Packages', level=2)
    doc.add_paragraph('requirements.txt:')
    
    req_code = doc.add_paragraph()
    req_code.style.font.name = 'Calibri'
    req_code.style.font.size = Pt(9)
    req_code.add_run('pandas>=2.0.0\n')
    req_code.add_run('openpyxl>=3.1.0\n')
    req_code.add_run('selenium>=4.15.0\n')
    req_code.add_run('requests>=2.31.0\n')
    req_code.add_run('webdriver-manager>=4.0.0\n')
    req_code.add_run('beautifulsoup4>=4.12.0')
    
    doc.add_heading('8.2 Package Purposes', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Package'
    hdr_cells[1].text = 'Purpose'
    hdr_cells[2].text = 'Usage'
    
    packages = [
        ('pandas', 'Excel file reading', 'pd.read_excel()'),
        ('openpyxl', 'Excel file writing', "df.to_excel(engine='openpyxl')"),
        ('selenium', 'Browser automation', 'WebDriver for fallback'),
        ('requests', 'HTTP client', 'Direct PDF downloads'),
        ('webdriver-manager', 'ChromeDriver management', 'Automatic driver setup'),
        ('beautifulsoup4', 'HTML parsing', 'Metadata extraction')
    ]
    
    for pkg, purpose, usage in packages:
        row_cells = table.add_row().cells
        row_cells[0].text = pkg
        row_cells[1].text = purpose
        row_cells[2].text = usage
    
    # Section 9: File Structure
    doc.add_page_break()
    doc.add_heading('9. File Structure', level=1)
    
    doc.add_heading('9.1 Source Files', level=2)
    
    doc.add_heading('patent_downloader_gui.py (940 lines)', level=3)
    doc.add_paragraph('Main application file, contains PatentDownloaderGUI class, all GUI and download logic')
    
    doc.add_heading('check_setup.py (102 lines)', level=3)
    doc.add_paragraph('Setup verification utility, checks Python, packages, Chrome, standalone diagnostic tool')
    
    doc.add_heading('9.2 Configuration Files', level=2)
    config_files = [
        ('requirements.txt', 'Python package dependencies, version specifications'),
        ('install_requirements.bat', 'Windows batch script, installs Python packages'),
        ('🚀 START HERE - GUI.bat', 'Windows launcher script, runs Python application')
    ]
    
    for filename, desc in config_files:
        para = doc.add_paragraph()
        run1 = para.add_run(f'{filename}: ')
        run1.bold = True
        para.add_run(desc)
    
    # Section 10: Implementation Details
    doc.add_page_break()
    doc.add_heading('10. Implementation Details', level=1)
    
    doc.add_heading('10.1 Threading Implementation', level=2)
    
    thread_code = doc.add_paragraph()
    thread_code.style.font.name = 'Calibri'
    thread_code.style.font.size = Pt(9)
    thread_code.add_run('download_thread = threading.Thread(target=self.download_patents, daemon=True)\n')
    thread_code.add_run('download_thread.start()')
    
    doc.add_paragraph('Thread Safety: GUI updates via root.update_idletasks() from download thread, no shared mutable state conflicts, is_downloading flag for graceful termination')
    
    doc.add_heading('10.2 Download Strategy', level=2)
    doc.add_paragraph('Priority Order:')
    priorities = [
        'Direct download from Google Patents (fastest)',
        'FreePatentsOnline fallback',
        'Browser method (if implemented, currently not primary)'
    ]
    for i, priority in enumerate(priorities, 1):
        doc.add_paragraph(f'{i}. {priority}', style='List Number')
    
    # Section 11: Error Handling Architecture
    doc.add_page_break()
    doc.add_heading('11. Error Handling Architecture', level=1)
    
    doc.add_heading('11.1 Exception Hierarchy', level=2)
    
    doc.add_heading('File Errors:', level=3)
    doc.add_paragraph('FileNotFoundError: Excel file not found')
    doc.add_paragraph('pd.errors: Pandas Excel reading errors')
    doc.add_paragraph('Handled with user-friendly messages')
    
    doc.add_heading('Network Errors:', level=3)
    doc.add_paragraph('requests.exceptions.RequestException: HTTP errors')
    doc.add_paragraph('requests.exceptions.Timeout: Timeout errors')
    doc.add_paragraph('Handled with fallback methods')
    
    doc.add_heading('11.2 Error Handling Strategy', level=2)
    doc.add_paragraph('Defensive Programming: Try-except blocks around critical operations, validation before operations, graceful degradation')
    doc.add_paragraph('Error Propagation: Critical errors logged and displayed, non-critical errors logged but don\'t stop process, user informed of failures')
    doc.add_paragraph('Recovery Mechanisms: Automatic fallback to alternative methods, failed patents logged for retry, process continues with next patent')
    
    # Section 12: Logging System
    doc.add_page_break()
    doc.add_heading('12. Logging System', level=1)
    
    doc.add_heading('12.1 Logging Architecture', level=2)
    doc.add_paragraph('Dual Logger System:')
    doc.add_paragraph('1. Main Logger: All operations and errors')
    doc.add_paragraph('2. Failed Patents Logger: Only failed downloads')
    
    doc.add_heading('12.2 Log Formats', level=2)
    
    doc.add_heading('Main Log Format:', level=3)
    log_format = doc.add_paragraph()
    log_format.style.font.name = 'Calibri'
    log_format.style.font.size = Pt(9)
    log_format.add_run('%(asctime)s - %(levelname)s - %(message)s\n')
    log_format.add_run('Example: 2025-12-15 10:30:45 - INFO - Downloading: US1234567A')
    
    doc.add_heading('Failed Patents Log Format:', level=3)
    failed_format = doc.add_paragraph()
    failed_format.style.font.name = 'Calibri'
    failed_format.style.font.size = Pt(9)
    failed_format.add_run('%(asctime)s - %(message)s\n')
    failed_format.add_run('Example: 2025-12-15 10:35:22 - FAILED | Original: US9999999A | ...')
    
    # Section 13: Testing Considerations
    doc.add_page_break()
    doc.add_heading('13. Testing Considerations', level=1)
    
    doc.add_heading('13.1 Unit Testing', level=2)
    doc.add_paragraph('Testable Components: clean_patent_number(), extract_patent_info(), read_patent_numbers(), construct_pdf_url()')
    doc.add_paragraph('Mock Requirements: Mock HTTP requests, mock file system operations, mock Excel files')
    
    doc.add_heading('13.2 Integration Testing', level=2)
    doc.add_paragraph('Test Scenarios: End-to-end download process, Excel file reading, PDF file saving, log file generation, Excel report creation')
    doc.add_paragraph('Test Data: Sample Excel files, known patent numbers, mock HTTP responses')
    
    # Section 14: Deployment
    doc.add_page_break()
    doc.add_heading('14. Deployment', level=1)
    
    doc.add_heading('14.1 Installation Process', level=2)
    
    doc.add_heading('Step 1: Prerequisites', level=3)
    prereqs = [
        'Install Python 3.8+',
        'Install Google Chrome',
        'Verify Python in PATH'
    ]
    for prereq in prereqs:
        doc.add_paragraph(prereq, style='List Number')
    
    doc.add_heading('Step 2: Package Installation', level=3)
    install_code = doc.add_paragraph()
    install_code.style.font.name = 'Calibri'
    install_code.add_run('pip install -r requirements.txt')
    
    doc.add_heading('Step 3: Verification', level=3)
    verify_code = doc.add_paragraph()
    verify_code.style.font.name = 'Calibri'
    verify_code.add_run('python check_setup.py')
    
    # Section 15: Performance Optimization
    doc.add_page_break()
    doc.add_heading('15. Performance Optimization', level=1)
    
    doc.add_heading('15.1 Current Optimizations', level=2)
    optimizations = [
        ('Direct Download Mode', 'Avoids browser overhead, 3-5x faster than browser method, lower memory usage'),
        ('Streaming Downloads', 'Chunked file writing, prevents memory issues with large PDFs, efficient for batch processing'),
        ('Threading', 'Non-blocking GUI during downloads, responsive user interface, can stop downloads mid-process')
    ]
    
    for opt_title, opt_desc in optimizations:
        para = doc.add_paragraph()
        run1 = para.add_run(f'{opt_title}: ')
        run1.bold = True
        para.add_run(opt_desc)
    
    doc.add_heading('15.2 Potential Optimizations', level=2)
    potential = [
        'Parallel Downloads: Multiple concurrent downloads, thread pool for parallel processing, rate limiting to avoid blocking',
        'Caching: Cache patent HTML pages, avoid re-fetching same patents, reduce network requests',
        'Batch Operations: Batch Excel updates, reduce file I/O operations, improve write performance',
        'Connection Pooling: Reuse HTTP connections, reduce connection overhead, faster subsequent requests'
    ]
    
    for pot in potential:
        para = doc.add_paragraph()
        parts = pot.split(':')
        run1 = para.add_run(parts[0] + ': ')
        run1.bold = True
        if len(parts) > 1:
            para.add_run(parts[1])

    # Section 16: System Recommendations
    doc.add_page_break()
    doc.add_heading('16. System Recommendations', level=1)

    doc.add_heading('16.1 Architecture Improvements', level=2)
    tech_recs = [
        ('Asynchronous I/O', 'Migrate from threading+requests to asyncio+aiohttp involved in networking for higher concurrency and non-blocking I/O.'),
        ('Database Integration', 'Implement SQLite/SQLAlchemy for robust state tracking, resuming, and sophisticated reporting capabilities.'),
        ('Dockerization', 'Containerize the application to ensure consistent runtime environments and simplify dependency management across different machines.')
    ]
    for title, desc in tech_recs:
        para = doc.add_paragraph()
        run1 = para.add_run(f'{title}: ')
        run1.bold = True
        para.add_run(desc)

    doc.add_heading('16.2 Maintainability', level=2)
    maint_recs = [
        'Implement Type Hinting (mypy) throughout the codebase for better static analysis',
        'Add comprehensive Unit Tests (pytest) with at least 80% coverage',
        'Set up a CI/CD pipeline (GitHub Actions) for automated testing and linting'
    ]
    for item in maint_recs:
        doc.add_paragraph(item, style='List Bullet')

    # Appendices
    doc.add_page_break()
    doc.add_heading('Appendix A: Code Examples', level=1)
    
    doc.add_heading('A.1 Direct Download Implementation', level=2)
    example_code = doc.add_paragraph()
    example_code.style.font.name = 'Calibri'
    example_code.style.font.size = Pt(9)
    example_code.add_run('def try_direct_download(self, patent_number):\n')
    example_code.add_run('    clean_number = self.clean_patent_number(patent_number)\n')
    example_code.add_run('    patent_url = f"https://patents.google.com/patent/{clean_number}"\n')
    example_code.add_run('    \n')
    example_code.add_run('    headers = {\n')
    example_code.add_run("        'User-Agent': 'Mozilla/5.0 ...'\n")
    example_code.add_run('    }\n')
    example_code.add_run('    \n')
    example_code.add_run('    response = requests.get(patent_url, headers=headers, timeout=10)\n')
    example_code.add_run('    response.raise_for_status()\n')
    example_code.add_run('    \n')
    example_code.add_run('    pdf_pattern = r\'https://patentimages\\.storage\\.googleapis\\.com/[^"\\\']+\\.pdf\'\n')
    example_code.add_run('    pdf_matches = re.findall(pdf_pattern, response.text)\n')
    example_code.add_run('    \n')
    example_code.add_run('    if pdf_matches:\n')
    example_code.add_run('        pdf_url = pdf_matches[0]\n')
    example_code.add_run('        if self.download_pdf_direct(pdf_url, clean_number):\n')
    example_code.add_run('            return True\n')
    example_code.add_run('    return False')
    
    doc.add_page_break()
    doc.add_heading('Appendix B: External APIs', level=1)
    
    doc.add_heading('B.1 Google Patents', level=2)
    doc.add_paragraph('Base URL: https://patents.google.com/patent/{patent_number}')
    doc.add_paragraph('PDF URL Pattern: https://patentimages.storage.googleapis.com/{hash}/{patent_number}.pdf')
    doc.add_paragraph('Rate Limiting: No official rate limits documented, 2-second delay implemented as precaution')
    
    doc.add_heading('B.2 FreePatentsOnline', level=2)
    doc.add_paragraph('PDF URL Pattern: https://www.freepatentsonline.com/{patent_number}.pdf')
    doc.add_paragraph('Availability: Not all patents available, used as fallback only')
    
    doc.add_page_break()
    doc.add_heading('Appendix C: Known Limitations', level=1)
    limitations = [
        'PDF URL Extraction: Relies on HTML structure, may break if Google changes format',
        'Patent Availability: Not all patents available on Google Patents',
        'Rate Limiting: No official rate limit information',
        'Browser Dependency: Requires Chrome for fallback method',
        'Windows Focus: Optimized for Windows, may need adjustments for other OS',
        'Excel Format: Requires specific column name ("Display Key")',
        'Single Thread: Downloads sequential, not parallel'
    ]
    
    for limitation in limitations:
        doc.add_paragraph(limitation, style='List Bullet')
    
    # Footer
    doc.add_paragraph()
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run('End of Technical Documentation')
    run.bold = True
    run.font.size = Pt(12)
    
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run('Google Patent PDF Downloader - Version 1.0')
    
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run('Document Version 1.0 - December 2025')
    
    # Save document
    doc.save('TECHNICAL_DOCUMENTATION.docx')
    print("Technical Documentation created: TECHNICAL_DOCUMENTATION.docx")

if __name__ == '__main__':
    print("Generating documentation files...")
    print()
    create_functional_documentation()
    create_technical_documentation()
    print()
    print("Documentation generation complete!")
    print("Files created:")
    print("  - FUNCTIONAL_DOCUMENTATION.docx")
    print("  - TECHNICAL_DOCUMENTATION.docx")
